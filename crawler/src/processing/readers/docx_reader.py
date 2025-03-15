"""
DOCX Document Reader

This module provides the DocxReader class for reading and extracting content from
Microsoft Word (.docx) files.
"""

import os
from typing import Any, Optional

from .base import DocumentReader
from ..document_content import DocumentContent


class DocxReader(DocumentReader):
    """Reader for Microsoft Word (.docx) files."""
    
    def read(self, file_path: str) -> DocumentContent:
        """Read DOCX file and extract content.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            DocumentContent object containing the extracted content
        """
        content = DocumentContent()
        try:
            import docx
            
            # Open the document
            doc = docx.Document(file_path)
            
            # Extract metadata
            doc_properties = doc.core_properties
            metadata = {
                'source': file_path,
                'format': 'docx',
                'title': doc_properties.title or "",
                'author': doc_properties.author or "",
                'subject': doc_properties.subject or "",
                'keywords': doc_properties.keywords or "",
                'created': str(doc_properties.created) if doc_properties.created else "",
                'modified': str(doc_properties.modified) if doc_properties.modified else "",
                'last_modified_by': doc_properties.last_modified_by or "",
                'size_bytes': os.path.getsize(file_path)
            }
            content.set_metadata(metadata)
            
            # Extract all paragraphs
            for para in doc.paragraphs:
                if para.text.strip():  # Only add non-empty paragraphs
                    content.add_text(para.text)
            
            # Extract tables
            table_index = 0
            for table in doc.tables:
                table_data = []
                for i, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        # Get text from each cell, handling potential nested content
                        cell_text = ""
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                cell_text += paragraph.text + " "
                        row_data.append(cell_text.strip())
                    table_data.append(" | ".join(row_data))
                
                table_text = "\n".join(table_data)
                content.add_table(table_text, None, table_index)
                table_index += 1
            
            # Handle images
            img_index = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_part = rel.target_part
                        image_data = image_part.blob
                        
                        # Process the image
                        description = self.process_image(image_data)
                        content.add_image(image_data, description, None, img_index)
                        img_index += 1
                    except Exception as img_err:
                        print(f"Error processing image in DOCX: {img_err}")
            
            return content
            
        except Exception as e:
            print(f"Error reading DOCX file {file_path}: {e}")
            content.add_text(f"Error reading file: {str(e)}")
            return content 